from __future__ import annotations

import csv
import os
import tempfile
import uuid
from pathlib import Path
from typing import Any
from urllib.request import Request, urlopen

import fitz
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from openpyxl import load_workbook
from pypdf import PdfReader

from src.agents.base import AgentRequest, AgentResponse, AgentSpec
from src.config import settings
from src.core.context import RuntimeContext
from src.core.oss import upload_to_oss
from src.core.runtime import runtime
from src.core.tool_prompt import render_agent_tool_prompt
from src.models.llm import chat_complete, vision_describe_image
from src.prompts.system import FILE_ANALYSIS_AGENT_PROMPT


PDF_MIME_TYPES = {"application/pdf", "application/x-pdf"}
WORD_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
EXCEL_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
TEXT_MIME_PREFIXES = ("text/",)
TEXT_MIME_TYPES = {
    "application/json",
    "application/xml",
    "application/yaml",
    "application/x-yaml",
    "application/javascript",
}
TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".json",
    ".csv",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".css",
    ".js",
    ".ts",
    ".tsx",
    ".vue",
    ".py",
    ".java",
    ".c",
    ".cc",
    ".cpp",
    ".h",
    ".hpp",
    ".go",
    ".rs",
    ".php",
    ".sql",
    ".sh",
    ".log",
    ".ini",
    ".conf",
    ".properties",
}
MAX_EXTRACT_CHARS = 24000
MAX_SUMMARY_INPUT_CHARS = 18000
MAX_EMBEDDED_IMAGES = 6
MAX_IMAGE_BYTES = 8 * 1024 * 1024


class ExtractedImage(dict[str, Any]):
    pass


class FileAnalysisAgent:
    spec = AgentSpec(
        name="file_analysis_agent",
        description="用 Python 解析文件内容，并调用大模型总结给主 Agent",
        tools=("file.extract_file_info_tool", "vision.extract_image_info_tool"),
        model_tier="large",
        capabilities=("file_parse", "image_extract", "summarize"),
    )

    async def analyze(self, context: RuntimeContext, attachments: list[dict[str, Any]], user_input: str) -> str:
        file_attachments = [item for item in attachments if not self._is_image(item)]
        if not file_attachments:
            return ""

        parts = ["文件分析 Agent 结果："]
        for index, attachment in enumerate(file_attachments, start=1):
            result = await runtime.run_tool(
                context,
                "file.extract_file_info_tool",
                caller_agent=self.spec.name,
                attachment=attachment,
                user_input=user_input,
            )
            if result.success:
                parts.append(f"{index}. {result.summary or result.data.get('summary') or ''}")
            else:
                parts.append(f"{index}. 文件解析失败：{result.error or 'unknown error'}")
        return "\n\n".join(parts)

    async def run(self, request: AgentRequest) -> AgentResponse:
        attachments = request.payload.get("attachments") if request.payload else None
        content = await self.analyze(request.context, attachments or request.context.attachments, request.user_input)
        return AgentResponse(True, content=content)

    async def _analyze_one(self, context: RuntimeContext, attachment: dict[str, Any], user_input: str) -> str:
        name = str(attachment.get("name") or "未命名文件")
        url = str(attachment.get("url") or "")
        mime = str(attachment.get("mime") or "unknown")
        size = attachment.get("size")
        metadata = f"文件名：{name}\nMIME：{mime}\n大小：{size}\nURL：{url}"
        temp_path: Path | None = None

        try:
            temp_path = self._download_to_temp(url, self._suffix(attachment))
            extracted_text = self._extract_text(attachment, temp_path)
            image_summaries = await self._analyze_embedded_images(attachment, temp_path)
        except Exception as exc:
            return f"{metadata}\n状态：Python 文件解析失败：{exc}"
        finally:
            if temp_path is not None:
                self._delete_temp_file(temp_path)

        if not extracted_text and not image_summaries:
            return f"{metadata}\n状态：暂未从该文件中提取到可用文本或可识别图片。"

        try:
            summary = await self._summarize_with_model(context, metadata, extracted_text, image_summaries, user_input)
        except Exception as exc:
            summary = (
                f"模型总结失败：{exc}\n"
                f"Python 已提取到以下文本片段，可供主 Agent 继续使用：\n"
                f"{extracted_text[:6000]}\n\n"
                f"内嵌图片识别结果：\n{image_summaries[:6000]}"
            )
        return f"{metadata}\n状态：已完成 Python 解析和模型总结。\n总结：\n{summary}"

    async def _summarize_with_model(
        self,
        context: RuntimeContext,
        metadata: str,
        extracted_text: str,
        image_summaries: str,
        user_input: str,
    ) -> str:
        user_prompt = (
            f"运行时信息：platform={context.platform}，userId={context.user_id}，role={context.role}。\n\n"
            f"用户问题：{user_input or '用户上传了文件，请识别并总结文件内容。'}\n\n"
            f"文件元信息：\n{metadata}\n\n"
            f"Python 提取的文件文本如下：\n{extracted_text[:MAX_SUMMARY_INPUT_CHARS]}\n\n"
            f"文件内嵌图片的视觉模型识别结果如下：\n{image_summaries[:MAX_SUMMARY_INPUT_CHARS]}\n\n"
            "请用中文给主 Agent 输出结构化文件分析结果，包含：文件主题、关键事实、图片中的关键信息、"
            "与用户问题相关的信息、可直接引用的结论。"
        )
        system_prompt = f"{FILE_ANALYSIS_AGENT_PROMPT}\n\n{render_agent_tool_prompt(self.spec)}"
        return await chat_complete(system_prompt, user_prompt, tier="large")

    async def _analyze_embedded_images(self, attachment: dict[str, Any], path: Path) -> str:
        images = self._extract_embedded_images(attachment, path)
        if not images:
            return ""

        summaries: list[str] = []
        for index, image in enumerate(images[:MAX_EMBEDDED_IMAGES], start=1):
            try:
                image_url = await upload_to_oss(
                    f"embedded/{uuid.uuid4().hex}.{image.get('ext') or 'png'}",
                    image["content"],
                    image.get("mime") or "image/png",
                )
                description = await vision_describe_image(
                    image_url,
                    "请用中文识别这张来自 PDF/Word 文件内的图片，重点提取图片中的文字、表格、证书、截图、图表和与用户问题有关的信息。",
                )
                summaries.append(
                    f"图片{index}：来源={image.get('source')}，url={image_url}\n视觉识别：{description}"
                )
            except Exception as exc:
                summaries.append(f"图片{index}：识别失败：{exc}")
        return "\n\n".join(summaries)

    def _extract_embedded_images(self, attachment: dict[str, Any], path: Path) -> list[ExtractedImage]:
        mime = str(attachment.get("mime") or "")
        suffix = self._suffix(attachment)
        if mime in PDF_MIME_TYPES or suffix == ".pdf":
            return self._extract_pdf_images(path)
        if mime in WORD_MIME_TYPES or suffix == ".docx":
            return self._extract_docx_images(path)
        return []

    def _extract_pdf_images(self, path: Path) -> list[ExtractedImage]:
        images: list[ExtractedImage] = []
        document = fitz.open(str(path))
        try:
            for page_index in range(min(document.page_count, 20)):
                page = document.load_page(page_index)
                for image_index, image_info in enumerate(page.get_images(full=True), start=1):
                    if len(images) >= MAX_EMBEDDED_IMAGES:
                        return images
                    xref = image_info[0]
                    extracted = document.extract_image(xref)
                    image_content = extracted.get("image") or b""
                    if not image_content or len(image_content) > MAX_IMAGE_BYTES:
                        continue
                    ext = extracted.get("ext") or "png"
                    images.append(
                        ExtractedImage(
                            content=image_content,
                            ext=ext,
                            mime=f"image/{'jpeg' if ext == 'jpg' else ext}",
                            source=f"pdf-page-{page_index + 1}-image-{image_index}",
                        )
                    )
        finally:
            document.close()
        return images

    def _extract_docx_images(self, path: Path) -> list[ExtractedImage]:
        document = Document(str(path))
        images: list[ExtractedImage] = []
        for rel in document.part.rels.values():
            if len(images) >= MAX_EMBEDDED_IMAGES:
                break
            if rel.reltype != RT.IMAGE:
                continue
            image_content = rel.target_part.blob
            if not image_content or len(image_content) > MAX_IMAGE_BYTES:
                continue
            content_type = rel.target_part.content_type or "image/png"
            ext = content_type.rsplit("/", 1)[-1].replace("jpeg", "jpg")
            images.append(
                ExtractedImage(
                    content=image_content,
                    ext=ext,
                    mime=content_type,
                    source=f"docx-image-{len(images) + 1}",
                )
            )
        return images

    def _download_to_temp(self, url: str, suffix: str) -> Path:
        if not url.startswith(("http://", "https://")):
            raise ValueError("文件 URL 非法")
        max_bytes = settings.max_upload_size_mb * 1024 * 1024
        request = Request(url, headers={"User-Agent": "yaai-agent/1.0"})
        fd, raw_path = tempfile.mkstemp(prefix="yaai_agent_file_", suffix=suffix or ".bin")
        temp_path = Path(raw_path)
        total = 0
        try:
            with urlopen(request, timeout=30) as response:
                with os.fdopen(fd, "wb") as target:
                    while True:
                        chunk = response.read(1024 * 1024)
                        if not chunk:
                            break
                        total += len(chunk)
                        if total > max_bytes:
                            raise ValueError(f"文件超过处理上限 {settings.max_upload_size_mb}MB")
                        target.write(chunk)
            return temp_path
        except Exception:
            try:
                os.close(fd)
            except OSError:
                pass
            self._delete_temp_file(temp_path)
            raise

    def _delete_temp_file(self, path: Path) -> None:
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass

    def _extract_text(self, attachment: dict[str, Any], path: Path) -> str:
        mime = str(attachment.get("mime") or "")
        suffix = self._suffix(attachment)
        if mime in PDF_MIME_TYPES or suffix == ".pdf":
            return self._extract_pdf(path)
        if mime in WORD_MIME_TYPES or suffix == ".docx":
            return self._extract_docx(path)
        if mime in EXCEL_MIME_TYPES or suffix == ".xlsx":
            return self._extract_xlsx(path)
        if suffix == ".csv":
            return self._extract_csv(path)
        if mime.startswith(TEXT_MIME_PREFIXES) or mime in TEXT_MIME_TYPES or suffix in TEXT_SUFFIXES:
            return self._extract_plain_text(path)
        content = path.read_bytes()[: MAX_EXTRACT_CHARS * 4]
        decoded = content.decode("utf-8", errors="ignore").strip()
        if decoded and len(decoded) >= max(20, path.stat().st_size // 20):
            return decoded[:MAX_EXTRACT_CHARS]
        return ""

    def _extract_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages[:30]:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())
            if sum(len(item) for item in pages) >= MAX_EXTRACT_CHARS:
                break
        return "\n\n".join(pages).strip()[:MAX_EXTRACT_CHARS]

    def _extract_docx(self, path: Path) -> str:
        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip()[:MAX_EXTRACT_CHARS]

    def _extract_xlsx(self, path: Path) -> str:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in workbook.worksheets[:5]:
            lines.append(f"[Sheet] {sheet.title}")
            for row in sheet.iter_rows(max_row=80, values_only=True):
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if values:
                    lines.append(" | ".join(values))
                if sum(len(item) for item in lines) >= MAX_EXTRACT_CHARS:
                    break
        return "\n".join(lines).strip()[:MAX_EXTRACT_CHARS]

    def _extract_csv(self, path: Path) -> str:
        lines: list[str] = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as source:
            rows = csv.reader(source)
            for _, row in zip(range(120), rows):
                lines.append(" | ".join(cell.strip() for cell in row))
        return "\n".join(lines).strip()[:MAX_EXTRACT_CHARS]

    def _extract_plain_text(self, path: Path) -> str:
        with path.open("rb") as source:
            content = source.read(MAX_EXTRACT_CHARS * 4)
        return content.decode("utf-8", errors="ignore").strip()[:MAX_EXTRACT_CHARS]

    def _is_image(self, attachment: dict[str, Any]) -> bool:
        mime = str(attachment.get("mime") or "")
        return attachment.get("type") == "image" or mime.startswith("image/")

    def _suffix(self, attachment: dict[str, Any]) -> str:
        name = str(attachment.get("name") or "")
        return "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""
